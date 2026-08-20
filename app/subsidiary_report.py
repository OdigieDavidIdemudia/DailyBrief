import os
from docx import Document
from datetime import datetime

def populate_subsidiary_docx(json_data: dict, template_path: str, output_path: str, subsidiary_name: str, date_str: str):
    doc = Document(template_path)
    
    # Process Paragraphs for Metadata
    for p in doc.paragraphs:
        if "Date:" in p.text:
            dt = datetime.strptime(date_str, "%Y-%m-%d")
            suffix = "th" if 11 <= dt.day <= 13 else {1: "st", 2: "nd", 3: "rd"}.get(dt.day % 10, "th")
            formatted_date = f"{dt.day}{suffix} {dt.strftime('%B %Y')}"
            p.text = f"Date: {formatted_date}"
            
        if "Submitted By:" in p.text:
            p.text = "Submitted By: David Idemudia Odigie"
            
        if "Subsidiary:" in p.text:
            p.text = f"Subsidiary: {subsidiary_name}"

    summary = json_data.get("summary") or {}
    if not isinstance(summary, dict): summary = {}
    sections = json_data.get("sections") or {}
    if not isinstance(sections, dict): sections = {}
    
    # Table 0: Summary
    if len(doc.tables) > 0:
        t0 = doc.tables[0]
        if len(t0.rows) > 3:
            t0.rows[1].cells[1].text = str(summary.get("overallSecurityMonitoringStatus", "N/A"))
            t0.rows[2].cells[1].text = str(summary.get("criticalIssues", "N/A"))
            t0.rows[3].cells[1].text = str(summary.get("escalations", "N/A"))

    # Table 1: Cortex XDR
    if len(doc.tables) > 1:
        t1 = doc.tables[1]
        xdr = sections.get("cortexXDR") or {}
        if not isinstance(xdr, dict): xdr = {}
        if len(t1.rows) > 7:
            t1.rows[1].cells[1].text = str(xdr.get("brokerVMStatus", "N/A"))
            t1.rows[2].cells[1].text = str(xdr.get("connectedAgents", "N/A"))
            t1.rows[3].cells[1].text = str(xdr.get("disconnectedAgents", "N/A"))
            t1.rows[4].cells[1].text = str(xdr.get("connectionLost", "N/A"))
            t1.rows[5].cells[1].text = str(xdr.get("pendingIncidents", "N/A"))
            t1.rows[6].cells[1].text = str(xdr.get("otherObservations", "N/A"))
            t1.rows[7].cells[1].text = str(xdr.get("priorityActionRequired", "N/A"))

    # Table 2: SIEM
    if len(doc.tables) > 2:
        t2 = doc.tables[2]
        siem = sections.get("siem") or {}
        if not isinstance(siem, dict): siem = {}
        if len(t2.rows) > 6:
            t2.rows[1].cells[1].text = str(siem.get("reportingDevices", "N/A"))
            t2.rows[2].cells[1].text = str(siem.get("nonReportingDevices", "N/A"))
            t2.rows[3].cells[1].text = str(siem.get("pendingIncidents", "N/A"))
            t2.rows[4].cells[1].text = str(siem.get("logIngestionHealth", "N/A"))
            t2.rows[5].cells[1].text = str(siem.get("otherObservations", "N/A"))
            t2.rows[6].cells[1].text = str(siem.get("priorityActionRequired", "N/A"))

    # Table 3: NAC
    if len(doc.tables) > 3:
        t3 = doc.tables[3]
        nac = sections.get("nac") or {}
        if not isinstance(nac, dict): nac = {}
        if len(t3.rows) > 6:
            t3.rows[1].cells[1].text = str(nac.get("implementationStatus", "N/A"))
            t3.rows[2].cells[1].text = str(nac.get("connectedDevices", "N/A"))
            t3.rows[3].cells[1].text = str(nac.get("compliancePoliciesActive", "N/A"))
            t3.rows[4].cells[1].text = str(nac.get("nonCompliantDevices", "N/A"))
            t3.rows[5].cells[1].text = str(nac.get("otherObservations", "N/A"))
            t3.rows[6].cells[1].text = str(nac.get("priorityActionRequired", "N/A"))

    # Table 4: DLP
    if len(doc.tables) > 4:
        t4 = doc.tables[4]
        dlp = sections.get("dlp") or {}
        if not isinstance(dlp, dict): dlp = {}
        if len(t4.rows) > 6:
            t4.rows[1].cells[1].text = str(dlp.get("systemStatus", "N/A"))
            t4.rows[2].cells[1].text = str(dlp.get("activePolicies", "N/A"))
            t4.rows[3].cells[1].text = str(dlp.get("integratedAgents", "N/A"))
            t4.rows[4].cells[1].text = str(dlp.get("policyViolations", "N/A"))
            t4.rows[5].cells[1].text = str(dlp.get("otherObservations", "N/A"))
            t4.rows[6].cells[1].text = str(dlp.get("priorityActionRequired", "N/A"))

    # Table 5: Web Proxy
    if len(doc.tables) > 5:
        t5 = doc.tables[5]
        wp = sections.get("webProxy") or {}
        if not isinstance(wp, dict): wp = {}
        if len(t5.rows) > 6:
            t5.rows[1].cells[1].text = str(wp.get("systemStatus", "N/A"))
            t5.rows[2].cells[1].text = str(wp.get("devicesOnboarded", "N/A"))
            t5.rows[3].cells[1].text = str(wp.get("activeRulesPolicies", "N/A"))
            t5.rows[4].cells[1].text = str(wp.get("coverageIssues", "N/A"))
            t5.rows[5].cells[1].text = str(wp.get("otherObservations", "N/A"))
            t5.rows[6].cells[1].text = str(wp.get("priorityActionRequired", "N/A"))

    # Table 6: SHELT
    if len(doc.tables) > 6:
        t6 = doc.tables[6]
        shelt = sections.get("shelt") or {}
        if not isinstance(shelt, dict): shelt = {}
        if len(t6.rows) > 4:
            t6.rows[1].cells[1].text = str(shelt.get("totalPendingIssues", "N/A"))
            t6.rows[2].cells[1].text = str(shelt.get("criticalFindings", "N/A"))
            t6.rows[3].cells[1].text = str(shelt.get("currentHealthStatus", "N/A"))
            t6.rows[4].cells[1].text = str(shelt.get("priorityActionRequired", "N/A"))

    # Table 7: WAF
    if len(doc.tables) > 7:
        t7 = doc.tables[7]
        waf = sections.get("waf") or {}
        if not isinstance(waf, dict): waf = {}
        if len(t7.rows) > 4:
            t7.rows[1].cells[1].text = str(waf.get("numberOfWebsites", "N/A"))
            t7.rows[2].cells[1].text = str(waf.get("mode", "N/A"))
            t7.rows[3].cells[1].text = str(waf.get("otherObservations", "N/A"))
            t7.rows[4].cells[1].text = str(waf.get("priorityActionRequired", "N/A"))

    # Table 8: Key Risks
    if len(doc.tables) > 8:
        t8 = doc.tables[8]
        risks = json_data.get("keyRisksRequiringCISOAttention") or []
        
        # Clear existing rows except header
        for i in range(len(t8.rows) - 1, 0, -1):
            row = t8.rows[i]
            tbl = t8._tbl
            tr = row._tr
            tbl.remove(tr)
            
        # Add new rows
        if isinstance(risks, list):
            for r in risks:
                row_cells = t8.add_row().cells
                row_cells[0].text = str(r.get("risk", "N/A")) if isinstance(r, dict) else str(r)
                row_cells[1].text = str(r.get("impact", "N/A")) if isinstance(r, dict) else "N/A"
                row_cells[2].text = str(r.get("owner", "N/A")) if isinstance(r, dict) else "N/A"
        elif isinstance(risks, dict):
            row_cells = t8.add_row().cells
            row_cells[0].text = str(risks.get("risk", "N/A"))
            row_cells[1].text = str(risks.get("impact", "N/A"))
            row_cells[2].text = str(risks.get("owner", "N/A"))
        else:
            # It's a string or other type
            row_cells = t8.add_row().cells
            row_cells[0].text = str(risks)
            row_cells[1].text = "N/A"
            row_cells[2].text = "N/A"

    doc.save(output_path)
    return output_path
