import os
from docx import Document
from datetime import datetime

def populate_handover_docx(json_data: dict, template_path: str, output_path: str, is_update: bool, location: str, date_str: str, duration: str = ""):
    """
    Reads the docx template, injects data into the tables, and saves to output_path.
    """
    doc = Document(template_path)
    
    # Process Paragraphs for Metadata (Location, Date, Duration)
    for p in doc.paragraphs:
        if "Location:" in p.text:
            p.text = f"Location:       {location}"
        if "Date:" in p.text:
            # Format date beautifully e.g. 29th September 2025
            dt = datetime.strptime(date_str, "%Y-%m-%d")
            suffix = "th" if 11 <= dt.day <= 13 else {1: "st", 2: "nd", 3: "rd"}.get(dt.day % 10, "th")
            formatted_date = f"{dt.day}{suffix} {dt.strftime('%B %Y')}"
            p.text = f"Date:            {formatted_date}"
            
        if "I will be out of office" in p.text and duration:
            p.text = duration
            
        # Process Risks and Blockers into the bulleted list under "Important Notes / Risks"
        # We find the paragraph that says "Important Notes / Risks", then inject bullets after it
        # Actually, python-docx doesn't easily let us replace bullet points perfectly without style tracking.
        # But we can look for specific placeholders.
        
    # Process Tables
    # In both templates, Table 0 is the Tasks table. Table 1 is Meetings.
    # The user said leave Meetings blank.
    if len(doc.tables) > 0:
        task_table = doc.tables[0]
        
        # Clear existing rows except header
        # python-docx doesn't easily let you delete rows cleanly, so we'll just overwrite existing ones and add new ones.
        # Actually, let's keep the header row and remove the rest, then add rows.
        # Removing rows in python-docx:
        for i in range(len(task_table.rows) - 1, 0, -1):
            row = task_table.rows[i]
            # hack to delete row
            tbl = task_table._tbl
            tr = row._tr
            tbl.remove(tr)
            
        if is_update:
            # Update format: ['Projects/Tasks', 'Previous Status', 'Current Status']
            tasks = json_data.get("update_tasks", [])
            for t in tasks:
                row_cells = task_table.add_row().cells
                row_cells[0].text = str(t.get("project_task", ""))
                row_cells[1].text = str(t.get("previous_status", ""))
                row_cells[2].text = str(t.get("current_status", ""))
        else:
            # Handover format: ['Projects/Tasks', 'Current Status', 'Next actions', 'Contact Person/ Email subject', 'Assignee']
            tasks = json_data.get("handover_tasks", [])
            for t in tasks:
                row_cells = task_table.add_row().cells
                row_cells[0].text = str(t.get("project_task", ""))
                row_cells[1].text = str(t.get("current_status", ""))
                row_cells[2].text = str(t.get("next_actions", ""))
                row_cells[3].text = str(t.get("contact_person", ""))
                row_cells[4].text = str(t.get("assignee", ""))
                
    # Now for risks and blockers, let's append them after "Important Notes / Risks"
    # To keep it simple, let's just find "Important Notes / Risks" and insert text after it.
    for i, p in enumerate(doc.paragraphs):
        if "Important Notes / Risks" in p.text:
            risks = json_data.get("risks_and_blockers", [])
            # Delete existing bullet points below it (up to "Ongoing Meetings & Commitments")
            # We'll just insert new paragraphs right after it.
            insert_idx = i + 1
            for risk in reversed(risks):
                new_p = doc.paragraphs[insert_idx].insert_paragraph_before(f"- {risk}")
                # We could set style='List Bullet' if needed, but plain text with hyphen works
            break

    doc.save(output_path)
    return output_path
