"""
app/downtime_report.py
Generates a DOCX Downtime Incident Report by editing the original template.
"""

import os
from docx import Document

def _replace_cell_text(cell, label: str, value: str):
    """Clear cell and add bold label + plain text."""
    # Clear existing text
    cell.text = ""
    # Add new formatted text
    p = cell.paragraphs[0]
    run_label = p.add_run(label)
    run_label.bold = True
    
    # Split value by newlines so we can keep the formatting
    lines = str(value).split("\n")
    # Join with standard newline, python-docx converts this to <w:br/>
    p.add_run("\n".join(lines))

def build_downtime_docx(data: dict) -> str:
    """
    Build a Downtime Incident Report DOCX from `data` using the existing 
    template and return the absolute path to the generated temp file.
    """
    # Load the template (assuming it's in the root folder)
    template_path = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "SOC06072026006.docx"))
    doc = Document(template_path)
    table = doc.tables[0]

    # Row 1: Start Date | Start Time | Downtime ID | Duration
    _replace_cell_text(table.cell(1, 0), "Start Date:\n", data.get("start_date", ""))
    _replace_cell_text(table.cell(1, 1), "Start Time:\n", data.get("start_time", ""))
    _replace_cell_text(table.cell(1, 3), "Downtime ID:\n", data.get("downtime_id", ""))
    _replace_cell_text(table.cell(1, 6), "Downtime Duration:\n", data.get("duration", ""))

    # Row 2: Reported By | Position
    _replace_cell_text(table.cell(2, 0), "Reported By: ", data.get("reported_by", ""))
    _replace_cell_text(table.cell(2, 3), "Position: ", data.get("position", ""))

    # Row 3: System/Service Affected
    _replace_cell_text(table.cell(3, 0), "System/Service Affected: ", data.get("system_affected", ""))

    # Row 4: Severity Level
    _replace_cell_text(table.cell(4, 0), "Severity Level: ", data.get("severity", ""))

    # Row 5: Impact Summary
    _replace_cell_text(table.cell(5, 0), "Impact Summary:\n", data.get("impact_summary", ""))

    # Row 6: Detection & Notification | Root Cause
    _replace_cell_text(table.cell(6, 0), "Detection and Notification:\n", data.get("detection_and_notification", ""))
    _replace_cell_text(table.cell(6, 4), "Root Cause Analysis:\n", data.get("root_cause_analysis", ""))

    # Row 7: Mitigation & Recovery | Preventive Measures
    _replace_cell_text(table.cell(7, 0), "Mitigation and Recovery Actions:\n", data.get("mitigation_and_recovery", ""))
    _replace_cell_text(table.cell(7, 4), "Preventive Measures:\n", data.get("preventive_measures", ""))

    # Row 8: Internal Comm | External Comm | Resource
    _replace_cell_text(table.cell(8, 0), "Internal Communication:\n", data.get("internal_communication", ""))
    _replace_cell_text(table.cell(8, 2), "External Communication:\n", data.get("external_communication", ""))
    _replace_cell_text(table.cell(8, 5), "Resource: ", data.get("resource", "N/A"))

    # Row 9: End Date | End Time
    _replace_cell_text(table.cell(9, 0), "End Date: ", data.get("end_date", ""))
    _replace_cell_text(table.cell(9, 4), "End Time: ", data.get("end_time", ""))

    # Row 10: Signatures (Fixed as per requirement)
    _replace_cell_text(table.cell(10, 0), "Reviewed by: ", "Fatima Jinadu\nPosition: Unit Head, Security Monitoring\nDate:\nSignature:")
    _replace_cell_text(table.cell(10, 4), "Approved by: ", "Jacobsen Brai\nPosition: Group Head, ISMS\nDate:\nSignature:")

    # Save to temp file
    out_dir = "/tmp"
    os.makedirs(out_dir, exist_ok=True)
    filename = f"downtime_{data.get('downtime_id', 'report').replace('/', '_')}.docx"
    out_path = os.path.join(out_dir, filename)
    
    doc.save(out_path)
    return out_path
